# 2022-2-11 用于通过有道翻译获取对应单词的中文翻译
# https://dict.youdao.com/search?q=word&keyfrom=new-fanyi.smartResult

import requests as r
from bs4 import BeautifulSoup
import re
import docx
import os.path


class WordTrans:
    """用于转换单词与中文含义"""

    def __init__(self, words: list):
        self.words = words
        self.url = 'https://dict.youdao.com/search?q=word&keyfrom=new-fanyi.smartResult'

    def get_word(self) -> list:
        """从有道获取数据"""
        trans_words = []
        for word in self.words:
            response = r.get(self.url.replace('word', word))
            response.encoding = response.apparent_encoding
            html = BeautifulSoup(response.text, features="html.parser")
            content = html.select('#phrsListTab > div.trans-container > ul')
            # / html / body / div[1] / div[2] / div[1] / div[2] / div[2] / div[1] / div[2] / ul
            trans_words.append(content[0].get_text())
        return trans_words

    def out_limit(self, word_mean: list):
        """限制单词含义的输出量"""
        mean_dict = []
        for mean in word_mean:
            mean = mean.split('\n')
            fir_mean = mean[1].split('，')
            fir_mean = fir_mean[0:2]
            fir_mean[0].replace('n. ', '')
            mean_dict.append(fir_mean)
        return mean_dict


class DocxTrans:
    def __init__(self):
        pass

    def read_docx(self, document_path: str):
        """获取docx文件数据"""
        r_docx = [para.text.replace("\n", "")
                  for para in docx.Document(document_path).paragraphs
                  if para.text.replace("\n", "").replace(" ", "")
                  ]
        if re.search('Topic', r_docx[0]):
            del r_docx[0]
        return r_docx

    def dump_docx(self, mean_list: list, out_path: str, statue=0):
        """创建docx文件并写入，0表示仅输出含义，1表示输出原文-含义对（待完善"""
        if statue == 0:
            document = docx.Document()
            for mean in mean_list:
                document.add_parafraph(mean)
            document.save(out_path)

    def get_file_list(self, folder_path: str, keyword: str) -> list:
        """批量处理时获取某文件夹内匹配文件列表（暂时不支持文件夹下的文件夹"""
        path_dir = os.listdir(folder_path)
        files_path = []
        for file_dir in path_dir:
            if re.search(keyword, file_dir):
                files_path.append(folder_path + '\\' + file_dir)
        return files_path



if __name__ == '__main__':
    file_trans = DocxTrans()
    files_list = file_trans.get_file_list(
        'C:\\Users\\xzqng\\Documents\\WeChat Files\\huongxia\\FileStorage\\File\\2022-02',
        keyword='沸腾')
    output_path = 'C:\\Users\\xzqng\\Documents\\WeChat Files\\huongxia\\FileStorage\\File\\2022-02\\Topic'
    count = 6
    for file_path in files_list:
        word_list = file_trans.read_docx(file_path)
        trans = WordTrans(word_list)
        out_full = trans.get_word()
        out_list = trans.out_limit(out_full)
        file_trans.dump_docx(out_list, output_path + str(count) + '.docx')
        count += 1
